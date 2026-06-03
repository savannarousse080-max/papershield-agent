from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def build_report_docx_bytes(payload: dict[str, Any], choices: dict[str, str] | None = None) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("python-docx is required for Word report export. Install python-docx.") from exc

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)

    document.add_heading("PaperShield 审阅诊断报告", level=0)
    _add_basic_info(document, payload)
    _add_analysis_summary(document, payload.get("analysis_summary") or {})
    _add_metrics(document, payload.get("metrics") or {})
    _add_workflow(document, payload.get("workflow") or {})
    _add_final_text(document, _merged_final_text(payload, choices or {}))
    _add_paragraph_reviews(document, payload.get("paragraphs") or [], choices or {})
    _add_review_items(document, payload.get("review_items") or [])
    _add_compliance(document)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def write_report_docx(payload: dict[str, Any], output_path: Path, choices: dict[str, str] | None = None) -> None:
    output_path.write_bytes(build_report_docx_bytes(payload, choices))


def _add_basic_info(document, payload: dict[str, Any]) -> None:
    document.add_heading("基本信息", level=1)
    rows = [
        ("领域", _domain_label(payload.get("domain"))),
        ("提示词方案", str(payload.get("prompt_profile") or "default")),
        ("处理模式", "仅分析" if payload.get("analysis_only") else "润色与诊断"),
        ("正文段落数", str(payload.get("paragraph_count", 0))),
        ("重试次数", str(payload.get("retry_count", 0))),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_analysis_summary(document, summary: dict[str, Any]) -> None:
    if not summary:
        return
    document.add_heading("文章级分析", level=1)
    overview = str(summary.get("overview") or "").strip()
    if overview:
        document.add_paragraph(overview)
    for label, key in [("优势", "strengths"), ("问题", "issues"), ("建议", "suggestions")]:
        values = _list(summary.get(key))
        if values:
            document.add_heading(label, level=2)
            for value in values:
                document.add_paragraph(value, style="List Bullet")


def _add_metrics(document, metrics: dict[str, Any]) -> None:
    document.add_heading("组合指标", level=1)
    rows = [
        ("困惑度代理", _number(metrics.get("average_rewritten_perplexity"))),
        ("模板词减少率", _percent(metrics.get("average_template_reduction"))),
        ("引用保留率", _percent(metrics.get("average_citation_retention"))),
        ("困惑度变化", _percent(metrics.get("average_perplexity_change"), signed=True)),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value


def _add_workflow(document, workflow: dict[str, Any]) -> None:
    document.add_heading("工作流", level=1)
    document.add_paragraph(f"后端：{workflow.get('backend_label') or workflow.get('backend') or '未知'}")
    document.add_paragraph(f"路线：{workflow.get('route_label') or workflow.get('route') or '未知'}")
    document.add_paragraph(f"人工复核：{workflow.get('manual_review_label') or ('已触发' if workflow.get('manual_review_required') else '未触发')}")
    steps = workflow.get("steps") or []
    if steps:
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "序号"
        table.rows[0].cells[1].text = "步骤"
        table.rows[0].cells[2].text = "含义"
        for index, step in enumerate(steps, 1):
            cells = table.add_row().cells
            cells[0].text = str(index)
            cells[1].text = str(step.get("label") or step.get("id") or "")
            cells[2].text = str(step.get("description") or "")


def _add_final_text(document, final_text: str) -> None:
    document.add_heading("当前最终稿", level=1)
    if not final_text.strip():
        document.add_paragraph("暂无最终稿。")
        return
    for paragraph in final_text.split("\n\n"):
        if paragraph.strip():
            document.add_paragraph(paragraph.strip())


def _add_paragraph_reviews(document, paragraphs: list[dict[str, Any]], choices: dict[str, str]) -> None:
    document.add_heading("逐段审阅", level=1)
    if not paragraphs:
        document.add_paragraph("暂无正文段落。")
        return
    for paragraph in paragraphs:
        index = paragraph.get("index", "")
        document.add_heading(f"段落 {index}", level=2)
        choice = _choice_for(paragraph, choices)
        document.add_paragraph(f"当前选择：{'保留原文' if choice == 'original' else '采纳润色'}")
        document.add_paragraph(f"状态：{paragraph.get('status_label') or paragraph.get('status') or '未知'}")
        document.add_paragraph(f"建议：{paragraph.get('recommendation_label') or paragraph.get('recommendation') or '建议复核'}")
        flags = paragraph.get("risk_flag_labels") or paragraph.get("risk_flags") or []
        document.add_paragraph(f"风险标记：{'、'.join(flags) if flags else '无'}")
        selected = _selected_paragraph_text(paragraph, choices)
        if selected:
            document.add_paragraph("选用文本：")
            document.add_paragraph(selected)


def _add_review_items(document, items: list[str]) -> None:
    document.add_heading("人工复核清单", level=1)
    values = items or ["复核事实、术语、论点和引注位置是否与原文一致。"]
    for item in values:
        document.add_paragraph(str(item), style="List Bullet")


def _add_compliance(document) -> None:
    document.add_heading("合规提示", level=1)
    for item in [
        "指标为本地代理信号，不代表任何外部 AI 检测器结果。",
        "请人工复核事实、术语、论点和引注。",
        "仅用于你有权编辑的学术草稿。",
    ]:
        document.add_paragraph(item, style="List Bullet")


def _merged_final_text(payload: dict[str, Any], choices: dict[str, str]) -> str:
    paragraphs = payload.get("paragraphs") or []
    blocks = payload.get("document_blocks") or []
    if blocks:
        parts: list[str] = []
        for block in blocks:
            if block.get("kind") == "paragraph":
                paragraph = _find_paragraph(paragraphs, block.get("paragraph_report_index"))
                if paragraph:
                    parts.append(_selected_paragraph_text(paragraph, choices))
            else:
                parts.append(str(block.get("text") or ""))
        return "\n\n".join(part for part in parts if part)
    if paragraphs:
        return "\n\n".join(_selected_paragraph_text(paragraph, choices) for paragraph in paragraphs)
    return str(payload.get("final_text") or "")


def _find_paragraph(paragraphs: list[dict[str, Any]], index: Any) -> dict[str, Any] | None:
    for paragraph in paragraphs:
        if str(paragraph.get("index")) == str(index):
            return paragraph
    return None


def _selected_paragraph_text(paragraph: dict[str, Any], choices: dict[str, str]) -> str:
    choice = _choice_for(paragraph, choices)
    if choice == "original":
        return str(paragraph.get("original_text") or "")
    return str(paragraph.get("rewritten_text") or paragraph.get("original_text") or "")


def _choice_for(paragraph: dict[str, Any], choices: dict[str, str]) -> str:
    index = str(paragraph.get("index"))
    choice = choices.get(index) or choices.get(str(paragraph.get("paragraph_report_index"))) or ""
    if choice in {"original", "rewritten"}:
        return choice
    return "original" if paragraph.get("status") == "fallback" else "rewritten"


def _domain_label(domain: Any) -> str:
    labels = {"law": "法学", "economics": "经济学", "general": "通用社科领域"}
    return labels.get(str(domain), str(domain or "未指定"))


def _list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item) for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value]
    return []


def _number(value: Any) -> str:
    return f"{value:.2f}" if isinstance(value, (int, float)) else "-"


def _percent(value: Any, signed: bool = False) -> str:
    if not isinstance(value, (int, float)):
        return "-"
    sign = "+" if signed and value > 0 else ""
    return f"{sign}{value * 100:.1f}%"
