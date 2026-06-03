from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any


SUPPORTED_INPUT_FORMATS = {"txt", "docx"}


@dataclass
class InputDocument:
    path: Path
    text: str
    source_format: str
    warnings: list[str] = field(default_factory=list)


def resolve_input_format(path: Path, requested: str = "auto") -> str:
    if requested != "auto":
        if requested not in SUPPORTED_INPUT_FORMATS:
            raise ValueError(f"Unsupported input format: {requested}. Supported formats: txt, docx")
        return requested

    suffix = path.suffix.lower()
    if suffix == ".txt":
        return "txt"
    if suffix == ".docx":
        return "docx"
    raise ValueError(f"Unsupported input format: {suffix or '<none>'}. Supported formats: .txt, .docx")


def read_input_document(path: Path, requested_format: str = "auto") -> InputDocument:
    source_format = resolve_input_format(path, requested_format)
    if source_format == "txt":
        return InputDocument(path=path, text=path.read_text(encoding="utf-8"), source_format="txt")
    return _read_docx(path)


def read_docx_bytes(filename: str, content: bytes) -> InputDocument:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for .docx input. Install python-docx or use .txt input.") from exc

    return _input_from_docx_document(Path(filename), Document(BytesIO(content)))


def write_optimized_document(source: InputDocument, final_text: str, output_path: Path) -> None:
    if source.source_format == "txt":
        output_path.write_text(final_text, encoding="utf-8")
        return
    _write_docx(source.path, final_text, output_path)


def _read_docx(path: Path) -> InputDocument:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for .docx input. Install python-docx or use .txt input.") from exc

    return _input_from_docx_document(path, Document(path))


def _input_from_docx_document(path: Path, document: Any) -> InputDocument:
    paragraphs = [paragraph.text.strip() for paragraph in _iter_docx_paragraphs(document) if paragraph.text.strip()]
    warnings: list[str] = []
    if document.tables:
        warnings.append("DOCX contains tables; table cell text was included, but complex layout must be reviewed manually.")
    return InputDocument(path=path, text="\n\n".join(paragraphs), source_format="docx", warnings=warnings)


def _write_docx(source_path: Path, final_text: str, output_path: Path) -> None:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for .docx output. Install python-docx or use .txt input.") from exc

    source_document = Document(source_path)
    chunks = [chunk.strip() for chunk in final_text.split("\n\n") if chunk.strip()]
    target_paragraphs = [paragraph for paragraph in _iter_docx_paragraphs(source_document) if paragraph.text.strip()]

    if len(chunks) == len(target_paragraphs):
        for paragraph, chunk in zip(target_paragraphs, chunks):
            paragraph.clear()
            paragraph.add_run(chunk)
        source_document.save(output_path)
        return

    fallback = Document()
    for chunk in chunks:
        fallback.add_paragraph(chunk)
    fallback.save(output_path)


def _iter_docx_paragraphs(document: Any):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
