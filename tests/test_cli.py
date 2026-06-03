import os
import json
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]


class CliTests(unittest.TestCase):
    def run_cli(self, *args, env=None):
        merged_env = os.environ.copy()
        merged_env.pop("OPENAI_API_KEY", None)
        merged_env.pop("ANTHROPIC_API_KEY", None)
        merged_env.pop("PAPERSHIELD_API_KEY", None)
        if env:
            merged_env.update(env)
        return subprocess.run(
            [sys.executable, "main.py", *args],
            cwd=ROOT,
            env=merged_env,
            text=True,
            capture_output=True,
            check=False,
        )

    def test_list_domains(self):
        result = self.run_cli("list-domains")

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("law", result.stdout)
        self.assertIn("economics", result.stdout)
        self.assertIn("general", result.stdout)

    def test_missing_input_file_has_clear_error(self):
        result = self.run_cli("optimize", "missing.txt", "--domain", "law")

        self.assertNotEqual(result.returncode, 0)
        self.assertIn("Input file not found", result.stderr)

    def test_optimize_writes_text_and_report_with_mock_provider(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "draft.txt"
            source.write_text("此外，数据安全问题需要完善[1]。", encoding="utf-8")

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            optimized = tmp_path / "draft_optimized.txt"
            report = tmp_path / "draft_report.txt"
            self.assertTrue(optimized.exists())
            self.assertTrue(report.exists())
            self.assertIn("[1]", optimized.read_text(encoding="utf-8"))
            self.assertIn("合规提示", report.read_text(encoding="utf-8"))

    def test_doctor_reports_mock_provider_ready(self):
        result = self.run_cli("doctor", env={"PAPERSHIELD_LLM_PROVIDER": "mock"})

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("Provider: mock", result.stdout)
        self.assertIn("Model: mock", result.stdout)
        self.assertIn("Ready for local demo", result.stdout)

    def test_doctor_detects_missing_real_provider_key(self):
        result = self.run_cli("doctor", env={"PAPERSHIELD_LLM_PROVIDER": "openai"})

        self.assertEqual(result.returncode, 2)
        self.assertIn("Missing API key", result.stdout)

    def test_test_paragraph_json_outputs_structured_payload(self):
        result = self.run_cli(
            "test-paragraph",
            "此外，数据安全问题需要完善[1]。",
            "--domain",
            "law",
            "--json",
            env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
        )

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["domain"], "law")
        self.assertIn("[1]", payload["final_text"])
        self.assertEqual(payload["paragraphs"][0]["index"], 1)
        self.assertIn("review_items", payload)
        self.assertIn("workflow", payload)
        self.assertIn(payload["workflow"]["backend"], {"simple", "langgraph"})
        self.assertEqual(payload["workflow"]["nodes"][0], "parse_document")
        self.assertEqual(payload["workflow"]["nodes"][-1], "assemble_output")

    def test_optimize_report_format_json_writes_structured_report(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "draft.txt"
            source.write_text("此外，数据安全问题需要完善[1]。", encoding="utf-8")

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                "--report-format",
                "json",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            report = tmp_path / "draft_report.json"
            self.assertTrue(report.exists())
            payload = json.loads(report.read_text(encoding="utf-8"))
            self.assertEqual(payload["domain"], "law")
            self.assertIn("paragraphs", payload)
            self.assertFalse((tmp_path / "draft_report.txt").exists())

    def test_unknown_auto_input_format_has_clear_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "draft.md"
            source.write_text("正文", encoding="utf-8")

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 2)
            self.assertIn("Unsupported input format", result.stderr)

    def test_optimize_docx_round_trip_with_mock_provider(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "draft.docx"
            document = Document()
            document.add_heading("问题提出", level=1)
            document.add_paragraph("此外，数据安全问题需要完善[1]。")
            document.save(source)

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                "--input-format",
                "docx",
                "--report-format",
                "both",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            optimized_docx = tmp_path / "draft_optimized.docx"
            self.assertTrue(optimized_docx.exists())
            text = "\n".join(paragraph.text for paragraph in Document(optimized_docx).paragraphs)
            self.assertIn("[1]", text)
            self.assertNotIn("此外", text)
            self.assertIn("法益完整性", text)
            self.assertTrue((tmp_path / "draft_report.txt").exists())
            self.assertTrue((tmp_path / "draft_report.json").exists())
            self.assertTrue((tmp_path / "draft_review.txt").exists())

    def test_optimize_docx_rewrites_table_cell_text(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "table-draft.docx"
            document = Document()
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "此外，数据安全问题需要完善[1]。"
            document.save(source)

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                "--input-format",
                "docx",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            optimized_docx = tmp_path / "table-draft_optimized.docx"
            optimized = Document(optimized_docx)
            table_text = optimized.tables[0].cell(0, 0).text
            self.assertIn("[1]", table_text)
            self.assertNotIn("此外", table_text)
            self.assertIn("法益完整性", table_text)

    def test_optimize_docx_all_reports_writes_html_and_markdown_review(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "draft.docx"
            document = Document()
            document.add_heading("问题提出", level=1)
            document.add_paragraph("此外，数据安全问题需要完善[1]。")
            document.save(source)

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                "--input-format",
                "docx",
                "--report-format",
                "all",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            report_html = tmp_path / "draft_report.html"
            report_docx = tmp_path / "draft_report.docx"
            review_md = tmp_path / "draft_review.md"
            self.assertTrue(report_html.exists())
            self.assertTrue(report_docx.exists())
            self.assertTrue(review_md.exists())
            html = report_html.read_text(encoding="utf-8")
            self.assertIn("PaperShield 合规润色诊断报告", html)
            self.assertIn("状态=", html)
            self.assertNotIn("status=", html)
            docx_text = "\n".join(paragraph.text for paragraph in Document(report_docx).paragraphs)
            self.assertIn("PaperShield 审阅诊断报告", docx_text)
            self.assertIn("人工复核清单", review_md.read_text(encoding="utf-8"))

    def test_analysis_only_docx_uses_provider_and_preserves_text(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "analysis.docx"
            document = Document()
            document.add_heading("问题提出", level=1)
            document.add_paragraph("此外，数据安全问题需要完善[1]。")
            document.save(source)

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                "--input-format",
                "docx",
                "--analysis-only",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            optimized = Document(tmp_path / "analysis_optimized.docx")
            text = "\n".join(paragraph.text for paragraph in optimized.paragraphs)
            self.assertIn("此外，数据安全问题需要完善[1]。", text)
            self.assertNotIn("法益完整性", text)

    def test_docx_writer_preserves_heading_style_when_references_are_grouped(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            source = tmp_path / "references.docx"
            document = Document()
            document.add_heading("问题提出", level=1)
            document.add_paragraph("此外，数据安全问题需要完善[1]。")
            document.add_paragraph("参考文献")
            document.add_paragraph("[1] 张三：《数据法研究》，2021。")
            document.save(source)

            result = self.run_cli(
                "optimize",
                str(source),
                "--domain",
                "law",
                "--output",
                str(tmp_path),
                "--input-format",
                "docx",
                env={"PAPERSHIELD_LLM_PROVIDER": "mock"},
            )

            self.assertEqual(result.returncode, 0, result.stderr)
            optimized = Document(tmp_path / "references_optimized.docx")
            self.assertEqual(len(optimized.paragraphs), 4)
            self.assertEqual(optimized.paragraphs[0].style.name, "Heading 1")
            self.assertEqual(optimized.paragraphs[2].text, "参考文献")
            self.assertEqual(optimized.paragraphs[3].text, "[1] 张三：《数据法研究》，2021。")

    def test_eval_fixtures_passes_builtin_samples(self):
        result = self.run_cli("eval-fixtures")

        self.assertEqual(result.returncode, 0, result.stderr)
        self.assertIn("Fixtures passed", result.stdout)
        self.assertIn("law_citation_reference", result.stdout)

    def test_eval_fixtures_json_outputs_structured_summary(self):
        result = self.run_cli("eval-fixtures", "--json")

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertGreaterEqual(payload["total"], 6)
        self.assertEqual(payload["failed"], 0)
        self.assertIn("metrics", payload)
        self.assertGreaterEqual(payload["metrics"]["total_paragraphs"], payload["total"])
        self.assertGreaterEqual(payload["metrics"]["average_citation_retention"], 0.0)
        self.assertGreaterEqual(payload["metrics"]["review_required"], 0)
        self.assertGreaterEqual(payload["metrics"]["fallback_count"], 0)
        self.assertIn("workflow_backends", payload["metrics"])
        self.assertTrue(payload["metrics"]["workflow_backends"])
        result_ids = {item["id"] for item in payload["results"]}
        self.assertIn("law_semantic_addition_review", result_ids)
        self.assertIn("economics_provider_failure", result_ids)
        self.assertIn("general_analysis_only", result_ids)
        self.assertTrue(all(item["passed"] for item in payload["results"]))

    def test_workflow_info_json_outputs_graph_topology(self):
        result = self.run_cli("workflow-info", "--json", env={"PAPERSHIELD_LLM_PROVIDER": "mock"})

        self.assertEqual(result.returncode, 0, result.stderr)
        payload = json.loads(result.stdout)
        self.assertEqual(payload["nodes"][0], "parse_document")
        self.assertEqual(payload["nodes"][-1], "assemble_output")
        self.assertIn(payload["active_backend"], {"simple", "langgraph"})


if __name__ == "__main__":
    unittest.main()
