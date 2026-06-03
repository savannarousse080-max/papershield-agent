import io
import os
import tempfile
import unittest
import warnings
import urllib.error
from unittest.mock import patch

warnings.filterwarnings("ignore", message="Using `httpx` with `starlette.testclient` is deprecated.*")
from fastapi.testclient import TestClient  # noqa: E402

from web.app import _safe_error_message, create_app  # noqa: E402
from web.provider_settings import reset_provider_runtime_for_tests  # noqa: E402


class WebAppTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.TemporaryDirectory()
        self.addCleanup(self.tempdir.cleanup)
        self.config_path = os.path.join(self.tempdir.name, "provider.local.json")
        self.env_patch = patch.dict(os.environ, {"PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path}, clear=False)
        self.env_patch.start()
        self.addCleanup(self.env_patch.stop)
        reset_provider_runtime_for_tests()
        self.client = TestClient(create_app())

    def test_healthz_reports_runtime_status(self):
        with patch.dict(os.environ, {"PAPERSHIELD_PROMPT_PROFILE": "research_writing_zh_word_v1"}, clear=False):
            response = self.client.get("/healthz")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "ok")
        self.assertEqual(payload["version"], "1.11")
        self.assertEqual(payload["prompt_profile"], "research_writing_zh_word_v1")
        self.assertIn(payload["provider"], {"openai", "mock", "anthropic"})
        self.assertEqual(payload["compliance_mode"], "local-demo")
        self.assertIn("security", payload)
        self.assertIn("limits", payload["security"])
        self.assertIn("fastapi", payload["dependencies"])
        self.assertIn("python-docx", payload["dependencies"])

    def test_runtime_policy_reports_demo_security_limits(self):
        response = self.client.get("/api/runtime/policy")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertIn("limits", payload)
        self.assertIn("max_upload_bytes", payload["limits"])
        self.assertIn("provider_config_enabled", payload)
        self.assertIn("admin_token_required", payload)
        self.assertIn("provider_base_url_policy", payload)

    def test_workflow_topology_endpoint_reports_langgraph_nodes(self):
        response = self.client.get("/api/workflow/topology")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["nodes"][0], "parse_document")
        self.assertEqual(payload["nodes"][-1], "assemble_output")
        self.assertIn(payload["active_backend"], {"simple", "langgraph"})

    def test_security_headers_are_applied(self):
        response = self.client.get("/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.headers["x-content-type-options"], "nosniff")
        self.assertEqual(response.headers["x-frame-options"], "SAMEORIGIN")
        self.assertIn("frame-ancestors 'self'", response.headers["content-security-policy"])

    def test_provider_status_reports_safe_configuration_without_secrets(self):
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_LLM_PROVIDER": "openai",
            "PAPERSHIELD_LLM_MODEL": "demo-model",
            "OPENAI_API_KEY": "secret-value",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            response = self.client.get("/api/provider/status")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["provider"], "openai")
        self.assertEqual(payload["model"], "demo-model")
        self.assertTrue(payload["configured"])
        self.assertTrue(payload["api_key_present"])
        self.assertNotIn("secret-value", response.text)

    def test_provider_presets_include_mainland_and_global_models(self):
        response = self.client.get("/api/provider/presets")

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        preset_ids = {preset["id"] for group in payload["groups"] for preset in group["presets"]}
        for preset_id in [
            "deepseek",
            "qwen",
            "zhipu",
            "kimi",
            "baidu_qianfan",
            "tencent_hunyuan",
            "volcengine_ark",
            "minimax",
            "iflytek_spark",
            "baichuan",
            "openai",
            "anthropic",
            "gemini",
        ]:
            self.assertIn(preset_id, preset_ids)

    def test_provider_config_saves_non_secret_fields_and_never_echoes_key(self):
        response = self.client.post(
            "/api/provider/config",
            json={
                "preset_id": "gemini",
                "provider": "openai-compatible",
                "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
                "model": "gemini-2.5-flash",
                "api_key": "secret-value",
                "prompt_profile": "research_writing_zh_word_v1",
                "timeout": 120,
                "max_retries": 0,
            },
        )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(payload["preset_id"], "gemini")
        self.assertTrue(payload["api_key_present"])
        self.assertNotIn("secret-value", response.text)

        status = self.client.get("/api/provider/status").json()
        self.assertEqual(status["provider"], "openai-compatible")
        self.assertEqual(status["model"], "gemini-2.5-flash")
        self.assertEqual(status["prompt_profile"], "research_writing_zh_word_v1")
        self.assertTrue(status["base_url_configured"])
        self.assertTrue(status["api_key_present"])

        reset_provider_runtime_for_tests()
        reloaded = self.client.get("/api/provider/config").json()
        self.assertEqual(reloaded["preset_id"], "gemini")
        self.assertEqual(reloaded["model"], "gemini-2.5-flash")
        self.assertFalse(reloaded["api_key_present"])

    def test_provider_config_rejects_private_or_insecure_base_url(self):
        for base_url in ["http://api.example.com/v1", "https://127.0.0.1:9", "https://localhost:8000"]:
            response = self.client.post(
                "/api/provider/config",
                json={
                    "preset_id": "custom",
                    "provider": "openai-compatible",
                    "base_url": base_url,
                    "model": "demo",
                    "prompt_profile": "default",
                    "timeout": 30,
                    "max_retries": 0,
                },
            )

            self.assertEqual(response.status_code, 400, base_url)
            self.assertIn("Provider base URL", response.json()["detail"])

    def test_provider_config_requires_admin_token_when_configured(self):
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_ADMIN_TOKEN": "admin-secret",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            response = client.post(
                "/api/provider/config",
                json={
                    "preset_id": "mock",
                    "provider": "mock",
                    "base_url": "",
                    "model": "mock",
                    "prompt_profile": "default",
                },
            )

            self.assertEqual(response.status_code, 403)

            session_denied = client.post("/api/provider/session")
            self.assertEqual(session_denied.status_code, 403)

            authorized = client.post(
                "/api/provider/config",
                headers={"X-PaperShield-Admin-Token": "admin-secret"},
                json={
                    "preset_id": "mock",
                    "provider": "mock",
                    "base_url": "",
                    "model": "mock",
                    "prompt_profile": "default",
                },
            )
            session_authorized = client.post(
                "/api/provider/session",
                headers={"X-PaperShield-Admin-Token": "admin-secret"},
            )

        self.assertEqual(authorized.status_code, 200, authorized.text)
        self.assertEqual(session_authorized.status_code, 200, session_authorized.text)
        self.assertTrue(session_authorized.json()["authenticated"])

    def test_external_model_call_requires_admin_token_when_configured(self):
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_ADMIN_TOKEN": "admin-secret",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            denied = client.post(
                "/api/optimize",
                data={"text": "此外，数据安全问题需要完善[1]。", "domain": "law", "provider_mode": "configured"},
            )
            mock_allowed = client.post(
                "/api/optimize",
                data={"text": "此外，数据安全问题需要完善[1]。", "domain": "law", "provider_mode": "mock"},
            )

        self.assertEqual(denied.status_code, 403)
        self.assertIn("Admin token required", denied.json()["detail"])
        self.assertEqual(mock_allowed.status_code, 200, mock_allowed.text)

    def test_provider_config_can_be_disabled_for_public_demo(self):
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_PROVIDER_CONFIG_ENABLED": "0",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            response = client.post(
                "/api/provider/config",
                json={
                    "preset_id": "mock",
                    "provider": "mock",
                    "base_url": "",
                    "model": "mock",
                    "prompt_profile": "default",
                },
            )

        self.assertEqual(response.status_code, 403)

    def test_disabled_provider_config_ignores_persisted_local_provider(self):
        with open(self.config_path, "w", encoding="utf-8") as handle:
            handle.write(
                """
{
  "preset_id": "custom",
  "provider": "openai-compatible",
  "base_url": "https://provider.example.com/v1",
  "model": "stale-real-model",
  "prompt_profile": "research_writing_zh_word_v1",
  "timeout": 120,
  "max_retries": 0
}
""".strip()
            )
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_PROVIDER_CONFIG_ENABLED": "0",
            "PAPERSHIELD_LLM_PROVIDER": "mock",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            status = client.get("/api/provider/status").json()
            optimize_response = client.post(
                "/api/optimize",
                data={"text": "此外，数据安全问题需要完善[1]。", "domain": "law", "provider_mode": "configured"},
            )

        self.assertEqual(status["provider"], "mock")
        self.assertEqual(status["model"], "mock")
        self.assertFalse(status["provider_config_enabled"])
        self.assertEqual(optimize_response.status_code, 200, optimize_response.text)
        self.assertIn("[1]", optimize_response.json()["final_text"])

    def test_web_provider_config_takes_priority_over_environment(self):
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_LLM_PROVIDER": "openai",
            "PAPERSHIELD_LLM_MODEL": "env-model",
            "OPENAI_API_KEY": "env-secret",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            response = client.post(
                "/api/provider/config",
                json={
                    "preset_id": "qwen",
                    "provider": "openai-compatible",
                    "base_url": "https://dashscope.aliyuncs.com/compatible-mode/v1",
                    "model": "qwen-plus",
                    "prompt_profile": "default",
                    "timeout": 90,
                    "max_retries": 0,
                },
            )

            self.assertEqual(response.status_code, 200, response.text)
            status = client.get("/api/provider/status").json()

        self.assertEqual(status["provider"], "openai-compatible")
        self.assertEqual(status["model"], "qwen-plus")
        self.assertEqual(status["timeout"], 90)
        self.assertFalse(status["api_key_present"])
        self.assertNotIn("env-secret", response.text)

    def test_provider_check_configured_reports_missing_runtime_key(self):
        with patch.dict(os.environ, {"PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path}, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            client.post(
                "/api/provider/config",
                json={
                    "preset_id": "gemini",
                    "provider": "openai-compatible",
                    "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
                    "model": "gemini-2.5-flash",
                    "prompt_profile": "default",
                    "timeout": 30,
                    "max_retries": 0,
                },
            )
            response = client.post("/api/provider/check", data={"provider_mode": "configured"})

        self.assertEqual(response.status_code, 400)
        self.assertIn("Missing API key", response.json()["detail"])
        self.assertNotIn("gemini-key", response.text)

    def test_optimize_provider_failure_returns_top_level_provider_error(self):
        original = "此外，数据安全问题需要完善[1]。"
        with patch.dict(os.environ, {"PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path}, clear=True):
            reset_provider_runtime_for_tests()
            client = TestClient(create_app())
            client.post(
                "/api/provider/config",
                json={
                    "preset_id": "custom",
                    "provider": "openai-compatible",
                    "base_url": "https://provider.invalid/v1",
                    "model": "unreachable-model",
                    "api_key": "runtime-secret",
                    "prompt_profile": "default",
                    "timeout": 1,
                    "max_retries": 0,
                },
            )
            with patch("urllib.request.urlopen", side_effect=urllib.error.URLError("connection refused")):
                response = client.post(
                    "/api/optimize",
                    data={"text": original, "domain": "law", "provider_mode": "configured"},
                )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertTrue(payload["provider_error"]["failed"])
        self.assertTrue(payload["provider_error"]["all_fallback"])
        self.assertEqual(payload["provider_error"]["fallback_count"], payload["paragraph_count"])
        self.assertIn("model", payload["provider_error"]["message"].lower())
        self.assertNotIn("runtime-secret", response.text)

    def test_safe_error_message_redacts_secret_patterns(self):
        message = _safe_error_message(RuntimeError("Bearer sk-secret api_key=plain x-api-key: other"))

        self.assertNotIn("sk-secret", message)
        self.assertNotIn("plain", message)
        self.assertNotIn("other", message)
        self.assertIn("<hidden>", message)

    def test_provider_check_mock_returns_ready_without_api_key(self):
        response = self.client.post("/api/provider/check", data={"provider_mode": "mock"})

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["status"], "ready")
        self.assertEqual(payload["provider"], "mock")
        self.assertIn("本地演示模型", payload["message"])

    def test_provider_check_env_reports_missing_key_as_safe_error(self):
        env = {
            "PAPERSHIELD_PROVIDER_CONFIG_PATH": self.config_path,
            "PAPERSHIELD_LLM_PROVIDER": "openai",
        }
        with patch.dict(os.environ, env, clear=True):
            reset_provider_runtime_for_tests()
            response = self.client.post("/api/provider/check", data={"provider_mode": "env"})

        self.assertEqual(response.status_code, 400)
        self.assertIn("Missing API key", response.json()["detail"])
        self.assertNotIn("OPENAI_API_KEY=", response.text)

    def test_optimize_text_with_mock_provider(self):
        response = self.client.post(
            "/api/optimize",
            data={
                "text": "此外，数据安全问题需要完善[1]。",
                "domain": "law",
                "provider_mode": "mock",
            },
        )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertIn("[1]", payload["final_text"])
        self.assertEqual(payload["domain"], "law")
        self.assertEqual(payload["paragraphs"][0]["index"], 1)
        self.assertIn("status_label", payload["paragraphs"][0])
        self.assertIn("risk_flags", payload["paragraphs"][0])
        self.assertIn("risk_flag_labels", payload["paragraphs"][0])
        self.assertIn("review_items", payload)
        self.assertIn("recommendation", payload["paragraphs"][0])
        self.assertIn("recommendation_label", payload["paragraphs"][0])
        self.assertIn("diff_segments", payload["paragraphs"][0])
        self.assertIn("document_blocks", payload)
        self.assertEqual(payload["document_blocks"][0]["kind"], "paragraph")
        self.assertEqual(payload["document_blocks"][0]["paragraph_report_index"], 1)
        self.assertIn("workflow", payload)
        self.assertIn(payload["workflow"]["backend"], {"simple", "langgraph"})
        self.assertEqual(payload["workflow"]["nodes"][0], "parse_document")
        self.assertEqual(payload["workflow"]["nodes"][-1], "assemble_output")
        self.assertIn(payload["workflow"]["route"], {"quality_accepted", "manual_review_required"})
        self.assertIn("manual_review_required", payload["workflow"])
        self.assertIn("review_gate", payload["workflow"]["nodes"])

    def test_optimize_rejects_oversized_text(self):
        with patch.dict(os.environ, {"PAPERSHIELD_MAX_TEXT_CHARS": "10"}, clear=False):
            client = TestClient(create_app())
            response = client.post(
                "/api/optimize",
                data={"text": "这是一段超过限制的文本。", "domain": "law", "provider_mode": "mock"},
            )

        self.assertEqual(response.status_code, 413)
        self.assertIn("Text input is too large", response.json()["detail"])

    def test_optimize_rejects_too_many_paragraphs(self):
        with patch.dict(os.environ, {"PAPERSHIELD_MAX_PARAGRAPHS": "1"}, clear=False):
            client = TestClient(create_app())
            response = client.post(
                "/api/optimize",
                data={"text": "第一段。\n\n第二段。", "domain": "law", "provider_mode": "mock"},
            )

        self.assertEqual(response.status_code, 413)
        self.assertIn("Too many paragraphs", response.json()["detail"])

    def test_optimize_analysis_only_with_mock_provider_returns_article_summary(self):
        response = self.client.post(
            "/api/optimize",
            data={
                "text": "此外，数据安全问题需要完善[1]。",
                "domain": "law",
                "provider_mode": "mock",
                "analysis_only": "true",
            },
        )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertTrue(payload["analysis_only"])
        self.assertEqual(payload["paragraphs"][0]["status"], "analysis_only")
        self.assertIn("analysis_summary", payload)
        self.assertIn("overview", payload["analysis_summary"])
        self.assertTrue(payload["analysis_summary"]["issues"])
        self.assertIn("此外，数据安全问题需要完善[1]。", payload["final_text"])
        self.assertIn("steps", payload["workflow"])
        self.assertEqual(payload["workflow"]["steps"][0]["label"], "解析文档")

    def test_report_docx_endpoint_exports_word_diagnostic_report(self):
        optimize_response = self.client.post(
            "/api/optimize",
            data={"text": "此外，数据安全问题需要完善[1]。", "domain": "law", "provider_mode": "mock"},
        )
        self.assertEqual(optimize_response.status_code, 200, optimize_response.text)

        response = self.client.post(
            "/api/report/docx",
            json={"payload": optimize_response.json(), "choices": {"1": "rewritten"}},
        )

        self.assertEqual(response.status_code, 200, response.text)
        self.assertIn(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            response.headers["content-type"],
        )
        self.assertGreater(len(response.content), 1024)
        from docx import Document

        document = Document(io.BytesIO(response.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("PaperShield 审阅诊断报告", text)
        self.assertIn("工作流", text)

    def test_document_blocks_preserve_structure_for_frontend_merge(self):
        response = self.client.post(
            "/api/optimize",
            data={
                "text": "一、问题提出\n\n此外，数据安全问题需要完善[1]。\n\n图1 数据结构\n\n参考文献\n\n[1] 张三：《数据法研究》，2021。",
                "domain": "law",
                "provider_mode": "mock",
            },
        )

        self.assertEqual(response.status_code, 200, response.text)
        blocks = response.json()["document_blocks"]
        self.assertEqual([block["kind"] for block in blocks], ["heading", "paragraph", "figure", "references"])
        self.assertEqual(blocks[0]["text"], "一、问题提出")
        self.assertEqual(blocks[1]["paragraph_report_index"], 1)
        self.assertIsNone(blocks[1]["text"])
        self.assertIn("参考文献", blocks[3]["text"])

    def test_optimize_uploaded_txt_with_mock_provider(self):
        response = self.client.post(
            "/api/optimize",
            data={"domain": "law", "provider_mode": "mock"},
            files={"file": ("draft.txt", io.BytesIO("此外，数据安全问题需要完善[1]。".encode("utf-8")), "text/plain")},
        )

        self.assertEqual(response.status_code, 200, response.text)
        self.assertIn("[1]", response.json()["final_text"])

    def test_optimize_rejects_oversized_upload(self):
        with patch.dict(os.environ, {"PAPERSHIELD_MAX_UPLOAD_BYTES": "8"}, clear=False):
            client = TestClient(create_app())
            response = client.post(
                "/api/optimize",
                data={"domain": "law", "provider_mode": "mock"},
                files={"file": ("draft.txt", io.BytesIO("超过上传限制的文本".encode("utf-8")), "text/plain")},
            )

        self.assertEqual(response.status_code, 413)
        self.assertIn("Uploaded file is too large", response.json()["detail"])

    def test_optimize_uploaded_docx_with_mock_provider(self):
        from docx import Document

        document = Document()
        document.add_heading("问题提出", level=1)
        document.add_paragraph("此外，数据安全问题需要完善[1]。")
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = self.client.post(
            "/api/optimize",
            data={"domain": "law", "provider_mode": "mock"},
            files={
                "file": (
                    "draft.docx",
                    buffer,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        self.assertEqual(response.status_code, 200, response.text)
        payload = response.json()
        self.assertEqual(payload["source_format"], "docx")
        self.assertIn("[1]", payload["final_text"])
        self.assertIn("问题提出", payload["final_text"])

    def test_optimize_requires_text_or_file(self):
        response = self.client.post(
            "/api/optimize",
            data={"domain": "law", "provider_mode": "mock"},
        )

        self.assertEqual(response.status_code, 400)
        self.assertIn("text or .txt/.docx file", response.json()["detail"])

    def test_home_page_exposes_review_workspace(self):
        response = self.client.get("/")

        self.assertEqual(response.status_code, 200)
        html = response.text
        self.assertIn("workbench-grid", html)
        self.assertIn("control-rail", html)
        self.assertIn("artifact-rail", html)
        self.assertIn('id="workspace-resizer"', html)
        self.assertIn('role="separator"', html)
        self.assertIn('id="review-workbench"', html)
        self.assertIn('id="result-summary"', html)
        self.assertIn("原文", html)
        self.assertIn("润色稿", html)
        self.assertIn("最终稿", html)
        self.assertIn("复制最终稿", html)
        self.assertIn("导出 Markdown", html)
        self.assertIn("导出 HTML", html)
        self.assertIn("导出 Word", html)
        self.assertIn("PaperShield 学术稿件智能审阅平台", html)
        self.assertIn('rel="icon"', html)
        self.assertIn("/static/styles.css?v=1.11", html)
        self.assertIn("/static/app.js?v=1.11", html)
        self.assertIn("面向一般社科领域", html)
        self.assertIn("论证清晰度、表达自然度、术语一致性与引文可核查性", html)
        self.assertIn("一般社科", html)
        self.assertIn('data-artifact-tab="final"', html)
        self.assertIn('data-artifact-tab="evidence"', html)
        self.assertIn('data-artifact-tab="workflow"', html)
        self.assertIn('data-sample="law"', html)
        self.assertIn("本地演示模型", html)
        self.assertIn("外部模型", html)
        self.assertIn("保存设置", html)
        self.assertIn("测试连接", html)
        self.assertIn("连接状态", html)
        self.assertIn("上传 .txt/.docx", html)
        self.assertIn('id="workflow-trace"', html)
        self.assertIn('id="user-notice-modal"', html)
        self.assertIn("用户须知", html)
        self.assertIn("我已阅读并同意", html)

    def test_home_page_exposes_model_settings_workspace(self):
        response = self.client.get("/")

        self.assertEqual(response.status_code, 200)
        html = response.text
        for expected in [
            'id="model-settings"',
            'id="provider-config-form"',
            'id="provider-mode-setting"',
            'id="provider-preset"',
            'id="provider-base-url"',
            'id="provider-model"',
            'id="provider-api-key"',
            'id="provider-profile"',
            'id="provider-timeout"',
            'id="provider-max-retries"',
            'id="save-provider"',
            'id="clear-provider-key"',
            'id="provider-summary"',
            'id="request-estimate"',
            'id="provider-alert"',
            'id="provider-diagnostics"',
            'id="provider-admin-token"',
            'id="provider-login"',
            'id="provider-logout"',
            'id="provider-auth-status"',
        ]:
            self.assertIn(expected, html)
        self.assertIn("接口类型", html)
        self.assertIn("失败重试次数", html)
        self.assertIn("登录配置", html)
        self.assertIn("退出登录", html)
        self.assertIn("help-dot", html)

    def test_static_javascript_contains_review_actions(self):
        response = self.client.get("/static/app.js")

        self.assertEqual(response.status_code, 200)
        script = response.text
        self.assertIn("renderReviewWorkbench", script)
        self.assertIn("keep-original", script)
        self.assertIn("accept-rewrite", script)
        self.assertIn("copy-final", script)
        self.assertIn("download-markdown", script)
        self.assertIn("download-html", script)
        self.assertIn("download-word", script)
        self.assertIn("buildMarkdownReport", script)
        self.assertIn("buildHtmlReport", script)
        self.assertIn("downloadWordReport", script)
        self.assertIn("renderAnalysisSummary", script)
        self.assertIn("check-provider", script)
        self.assertIn("renderDiffSegments", script)
        self.assertIn("recommendation_label", script)
        self.assertIn("document_blocks", script)
        self.assertIn("analysis_only", script)
        self.assertIn("SAMPLE_DRAFTS", script)
        self.assertIn("activateArtifactTab", script)
        self.assertIn("renderResultSummary", script)
        self.assertIn("initializeWorkbenchResizer", script)
        self.assertIn("WORKBENCH_WIDTH_KEY", script)
        self.assertNotIn("merged.includes(paragraph.rewritten_text)", script)
        self.assertNotIn("merged.replace(paragraph.rewritten_text", script)

    def test_static_javascript_contains_model_settings_actions(self):
        response = self.client.get("/static/app.js")

        self.assertEqual(response.status_code, 200)
        script = response.text
        for expected in [
            "loadProviderSettings",
            "/api/provider/presets",
            "/api/provider/config",
            "/api/provider/session",
            "saveProviderConfig",
            "clearProviderKey",
            "loginProviderConfig",
            "logoutProviderConfig",
            "applyPresetToForm",
            "updateRequestEstimate",
            "provider_error",
            "initializeCustomSelects",
            "syncCustomSelect",
            "custom-select-source",
            "setProviderStatus",
            "X-PaperShield-Admin-Token",
            "sessionStorage",
            "模型调用失败，已保留原文",
        ]:
            self.assertIn(expected, script)

    def test_static_javascript_surfaces_runtime_policy(self):
        response = self.client.get("/static/app.js")

        self.assertEqual(response.status_code, 200)
        script = response.text
        self.assertIn("/api/runtime/policy", script)
        self.assertIn("applyRuntimePolicy", script)
        self.assertIn("provider_config_enabled", script)
        self.assertIn("admin_token_required", script)
        self.assertIn("公开演示模式已锁定模型配置", script)
        self.assertIn("请输入访问口令并登录", script)

    def test_static_javascript_contains_user_notice_modal_logic(self):
        response = self.client.get("/static/app.js")

        self.assertEqual(response.status_code, 200)
        script = response.text
        self.assertIn("initializeUserNotice", script)
        self.assertIn("showUserNotice", script)
        self.assertIn("acceptUserNotice", script)
        self.assertIn("USER_NOTICE_ACCEPTED_KEY", script)
        self.assertIn("localStorage", script)

    def test_static_javascript_contains_workflow_trace(self):
        response = self.client.get("/static/app.js")

        self.assertEqual(response.status_code, 200)
        script = response.text
        self.assertIn("renderWorkflowTrace", script)
        self.assertIn("renderWorkflowStep", script)
        self.assertIn("formatWorkflowRouteLabel", script)
        self.assertIn("workflow-trace", script)
        self.assertIn("manual_review_required", script)
        self.assertIn("解析文档", script)

    def test_static_javascript_localizes_review_enums(self):
        response = self.client.get("/static/app.js")

        self.assertEqual(response.status_code, 200)
        script = response.text
        self.assertIn("formatStatusLabel", script)
        self.assertIn("formatRiskFlagLabel", script)
        self.assertIn("formatChoiceLabel", script)
        self.assertIn("formatParagraphStatusLabel", script)
        self.assertIn("formatParagraphRiskFlagList", script)
        self.assertIn('"accepted": "已通过"', script)
        self.assertIn('"below_threshold": "低于质量阈值"', script)
        self.assertIn('"fallback": "已保留原文"', script)
        self.assertIn('"template_word_residue": "模板化表达残留"', script)
        self.assertIn('"citation_retention_risk": "引注保留风险"', script)
        self.assertNotIn("状态: ${paragraph.status}", script)
        self.assertNotIn("(paragraph.risk_flags || []).join", script)
        self.assertNotIn("|| status ||", script)
        self.assertNotIn("|| flag ||", script)


if __name__ == "__main__":
    unittest.main()
